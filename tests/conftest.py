import pytest
import os
import fitz
import docx
from pathlib import Path
from zipfile import ZipFile

@pytest.fixture
def temp_dir(tmp_path):
    return tmp_path

@pytest.fixture
def native_pdf(temp_dir):
    path = temp_dir / "native.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Heading 1", fontsize=20)
    page.insert_text((50, 100), "Some text here that is long enough to avoid the OCR trigger which is set at 50 characters. " * 2)
    doc.set_toc([[1, "Heading 1", 1]])
    doc.save(str(path))
    doc.close()
    return path

@pytest.fixture
def heuristic_pdf(temp_dir):
    path = temp_dir / "heuristic.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Large Heading", fontsize=18)
    page.insert_text((50, 100), "Normal text that is also very long, much longer than fifty characters, so it definitely does not trigger OCR logic.", fontsize=11)
    doc.save(str(path))
    doc.close()
    return path

@pytest.fixture
def password_pdf(temp_dir):
    path = temp_dir / "password.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Secret")
    doc.save(str(path), encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="secret")
    doc.close()
    return path

@pytest.fixture
def corrupt_pdf(temp_dir):
    path = temp_dir / "corrupt.pdf"
    path.write_bytes(b"%PDF-1.4\n%Corrupted data")
    return path

@pytest.fixture
def native_docx(temp_dir):
    path = temp_dir / "native.docx"
    doc = docx.Document()
    doc.add_paragraph('Main Title', style='Heading 1')
    doc.add_paragraph('Some content.')
    doc.add_paragraph('Subtitle', style='Heading 2')
    doc.save(str(path))
    return path

@pytest.fixture
def plain_txt(temp_dir):
    path = temp_dir / "plain.txt"
    path.write_text("Just some plain text without any structure.", encoding="utf-8")
    return path

@pytest.fixture
def latin1_txt(temp_dir):
    path = temp_dir / "latin1.txt"
    # Create some text that requires latin1
    path.write_bytes("Café and résumé".encode("latin-1"))
    return path

@pytest.fixture
def empty_file(temp_dir):
    path = temp_dir / "empty.txt"
    path.touch()
    return path

@pytest.fixture
def mixed_zip(temp_dir, plain_txt, native_docx):
    path = temp_dir / "mixed.zip"
    nested_zip_path = temp_dir / "nested.zip"
    
    with ZipFile(str(nested_zip_path), 'w') as z:
        z.write(str(plain_txt), plain_txt.name)
        
    with ZipFile(str(path), 'w') as z:
        z.write(str(native_docx), native_docx.name)
        z.write(str(nested_zip_path), nested_zip_path.name)
        
    return path
